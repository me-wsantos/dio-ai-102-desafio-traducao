import requests
from docx import Document
import os

subscription_key = "your subscription key"
endpoint = 'https://api.cognitive.microsofttranslator.com/'
location = "eastus2"
language_destination = 'pt-br'

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId': str(os.urandom(16))
  }

  body = [{
      'text': text
  }]

  params = {
      'api-version': '3.0',
      'from': 'en',
      'to': target_language
  }
  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

#textTranslated = translator_text("I think of you, I haven’t slept", language_destination)
#print(textTranslated)

def translate_document(path):
  document = Document(path)
  full_text = []
  
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)

  path_translated = path.replace(".docx", f"{language_destination}.docx")
  
  #Cria um novo documento com o conteúdo traduzido
  translated_doc.save(path_translated)
  return path_translated

input_file = "./fileToTranslate/music.docx"
translate_document(input_file)